from django.http import HttpResponse
from django.template.loader import render_to_string
from django.conf import settings
import os
import base64
from io import BytesIO

# Handle weasyprint import gracefully
try:
    from weasyprint import HTML, CSS
    WEASYPRINT_AVAILABLE = True
except Exception as e:
    WEASYPRINT_AVAILABLE = False

# Handle image processing imports
try:
    from PIL import Image, ImageDraw, ImageFont
    import fitz  # PyMuPDF
    IMAGE_PROCESSING_AVAILABLE = True
except ImportError as e:
    IMAGE_PROCESSING_AVAILABLE = False

# Handle office file processing imports
try:
    import docx
    import openpyxl
    OFFICE_PROCESSING_AVAILABLE = True
except ImportError:
    OFFICE_PROCESSING_AVAILABLE = False


def docx_to_html(file_content):
    """
    Converts the content of a .docx file to a simple HTML string.
    """
    if not OFFICE_PROCESSING_AVAILABLE:
        return "<html><body><p>DOCX processing library not available.</p></body></html>"
    try:
        document = docx.Document(BytesIO(file_content))
        html = "<html><head><style>body { font-family: sans-serif; } table { border-collapse: collapse; width: 100%; } td, th { border: 1px solid #dddddd; text-align: left; padding: 8px; }</style></head><body>"
        html += "<h1>Documento Word</h1>"
        for para in document.paragraphs:
            html += f"<p>{para.text}</p>"
        for table in document.tables:
            html += "<table>"
            for row in table.rows:
                html += "<tr>"
                for cell in row.cells:
                    html += f"<td>{cell.text}</td>"
                html += "</tr>"
            html += "</table><br>"
        html += "</body></html>"
        return html
    except Exception as e:
        return f"<html><body><p>Error converting DOCX: {e}</p></body></html>"

def xlsx_to_html(file_content):
    """
    Converts the content of a .xlsx file to an HTML table string.
    """
    if not OFFICE_PROCESSING_AVAILABLE:
        return "<html><body><p>XLSX processing library not available.</p></body></html>"
    try:
        workbook = openpyxl.load_workbook(BytesIO(file_content))
        sheet = workbook.active
        html = "<html><head><style>body { font-family: sans-serif; } table { border-collapse: collapse; width: 100%; } td, th { border: 1px solid #dddddd; text-align: left; padding: 8px; }</style></head><body>"
        html += "<h1>Planilha Excel</h1>"
        html += "<table>"
        for row in sheet.iter_rows():
            html += "<tr>"
            for cell in row:
                cell_value = cell.value if cell.value is not None else ""
                html += f"<td>{cell_value}</td>"
            html += "</tr>"
        html += "</table>"
        html += "</body></html>"
        return html
    except Exception as e:
        return f"<html><body><p>Error converting XLSX: {e}</p></body></html>"

def generate_pdf_response(template_name, context, css_path, filename):
    """
    Gera uma resposta HTTP com um PDF a partir de um template HTML.
    """
    if not WEASYPRINT_AVAILABLE:
        return HttpResponse("PDF generation is not available. WeasyPrint could not be loaded.", status=500)
    
    html_string = render_to_string(template_name, context)
    
    try:
        css_string = open(css_path, 'r').read()
        css = CSS(string=css_string)
    except FileNotFoundError:
        css = None

    try:
        html = HTML(string=html_string, base_url=settings.STATIC_ROOT)
        pdf_file = html.write_pdf(stylesheets=[css] if css else None)
        
        response = HttpResponse(pdf_file, content_type='application/pdf')
        response['Content-Disposition'] = f'attachment; filename="{filename}"'
        
        return response
    except Exception as e:
        return HttpResponse(f"Error generating PDF: {str(e)}", status=500)


def _create_placeholder_image(nome_arquivo, error_message):
    """Creates a placeholder image with an error message."""
    img_base64 = None
    try:
        try:
            font = ImageFont.truetype("arial.ttf", 15)
        except IOError:
            font = ImageFont.load_default()

        img = Image.new('RGB', (800, 100), color = (230, 230, 230))
        d = ImageDraw.Draw(img)
        d.text((10,10), f"Could not render: {nome_arquivo}", fill=(200,0,0), font=font)
        # Wrap error message text
        lines = []
        line = ""
        for word in str(error_message).split():
            if d.textlength(line + word) <= 780:
                line += word + " "
            else:
                lines.append(line)
                line = word + " "
        lines.append(line)

        y_text = 35
        for line in lines:
            d.text((10, y_text), line, fill=(0,0,0), font=font)
            y_text += 20

        buffer = BytesIO()
        img.save(buffer, format='PNG')
        img_base64 = base64.b64encode(buffer.getvalue()).decode('utf-8')
    except Exception as placeholder_error:
        print(f"Critical error creating placeholder image: {placeholder_error}")

    return img_base64

def process_attachments_for_pdf(attachments):
    """
    Processes various attachment types for inclusion in a PDF.
    """
    if not IMAGE_PROCESSING_AVAILABLE or not WEASYPRINT_AVAILABLE:
        return []
    
    processed_attachments = []
    
    for anexo in attachments:
        img_base64 = None
        nome_arquivo = "Unknown"
        descricao_anexo = ""
        try:
            file_field = getattr(anexo, 'arquivo', getattr(anexo, 'anexo', getattr(anexo, 'imagem', None)))
            if not file_field:
                continue

            nome_arquivo = getattr(anexo, 'nome_original', os.path.basename(file_field.name))
            descricao_anexo = getattr(anexo, 'descricao', '')
            
            file_field.seek(0)
            file_content = file_field.read()
            file_field.seek(0)
            
            file_extension = nome_arquivo.lower().split('.')[-1] if nome_arquivo else ''
            
            pdf_bytes_for_conversion = None

            if file_extension == 'pdf':
                pdf_bytes_for_conversion = file_content
            elif file_extension == 'docx':
                html_content = docx_to_html(file_content)
                pdf_bytes_for_conversion = HTML(string=html_content).write_pdf()
            elif file_extension == 'xlsx':
                html_content = xlsx_to_html(file_content)
                pdf_bytes_for_conversion = HTML(string=html_content).write_pdf()
            elif file_extension in ['jpg', 'jpeg', 'png', 'gif', 'bmp', 'webp']:
                img = Image.open(BytesIO(file_content))
                if img.mode in ('RGBA', 'P'):
                    img = img.convert('RGB')
                buffer = BytesIO()
                img.save(buffer, format='JPEG')
                img_base64 = base64.b64encode(buffer.getvalue()).decode('utf-8')

            if pdf_bytes_for_conversion:
                pdf_doc = fitz.open(stream=pdf_bytes_for_conversion, filetype="pdf")
                page = pdf_doc.load_page(0)
                pix = page.get_pixmap(dpi=150)
                img_bytes = pix.tobytes("png")
                img_base64 = base64.b64encode(img_bytes).decode('utf-8')

        except Exception as e:
            print(f"Error processing attachment ID {getattr(anexo, 'id', 'N/A')} ({nome_arquivo}): {e}")
            img_base64 = _create_placeholder_image(nome_arquivo, e)

        if img_base64:
            processed_attachments.append({
                'nome': nome_arquivo,
                'descricao': descricao_anexo,
                'is_image': True,
                'base64_data': img_base64
            })
    
    return processed_attachments